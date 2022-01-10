from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt

from ..utils import max_column_length


def convert(cv, stream):
    data = cv.data
    doc = Document()
    
    doc.core_properties.author = data['name']
    doc.core_properties.author = 'CV'
    doc.core_properties.comments = 'Generated via Python 3 and `docx`'
    doc.core_properties.created = datetime.utcnow()
    doc.core_properties.keywords = 'résumé cv'
    
    lb2_style = doc.styles['List Bullet 2']
    lb2_style.paragraph_format.left_indent = 0
    lb2_style.paragraph_format.space_after = 0

    bt_style = doc.styles['Body Text']
    bt_style.paragraph_format.space_after = 0

    sec = doc.sections[-1]
    sec.top_margin = Inches(.5)
    sec.bottom_margin = Inches(.5)
    sec.left_margin = Inches(.5)
    sec.right_margin = Inches(.5)
        
    # if image: {image}

    doc.add_heading(data['name'])
    urls = ' · '.join(data['urls'])
    doc.add_paragraph('{email} · {tel}\n{urls}'.format(
        email=data['email'],
        tel=data['tel'],
        urls=urls or ''
    ))
    
    doc.add_heading('Summary', level=2)
    doc.add_paragraph(data['summary'])
    
    doc.add_heading('Skills', level=2)
    mx = 1 + max_column_length(data['skills'], 'key')
    p = doc.add_paragraph()
    for i, skill in enumerate(data['skills']):
        p.add_run('{}{:{}}'.format(
            '\n' if i else '',
            '{}:'.format(skill['key']),
            mx
        )).bold = True
        p.add_run('\n{}'.format(skill['value']))

    doc.add_heading('Work Experience', level=2)
    for exp in data['experience']:
        p = doc.add_paragraph()
        p.style = bt_style
        r = p.add_run('{} '.format(exp['org']))
        r.bold = True
        r.font.size = Pt(13)
        p.add_run('· {}\n'.format(exp['location'])).italic = True

        for i,pos in enumerate(exp['positions']):
            if i:
                p = doc.add_paragraph()
                p.style = bt_style

            p.add_run('{} '.format(pos['title'])).bold = True
            p.add_run('· {} - {}'.format(
                pos['period']['from'],
                pos['period']['to']
            )).italic = True

            desc = pos.get('description')
            if desc:
                p.add_run('\n{}'.format(desc)).italic = True
            
            achs = pos.get('achievements', [])
            if achs:
                for ach in achs:
                    p = doc.add_paragraph(ach)
                    p.style = lb2_style

                p.style  = 'List Bullet'
            else:
                p.add_run('\n')

    if data['education']:
        doc.add_heading('Education', level=2)
        p = doc.add_paragraph()
        for i, edu in enumerate(data['education']):
            deg = edu.get('degree', '')
            dis = edu.get('descipline', '')
            p.add_run('{}{}, {}{}{}'.format(
                '\n' if i else '',
                edu['entity'],
                edu['location'],
                ', {}'.format(deg) if deg else '',
                ', {}'.format(dis) if dis else ''
            ))

    p = doc.add_paragraph()
    p.add_run('Last Modified: {}'.format(cv.date_updated)).italic = True
    doc.save(stream)

