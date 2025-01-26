from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt


def convert(cv, stream):
    doc = Document()

    doc.core_properties.author = cv.full_name
    doc.core_properties.author = "CV"
    doc.core_properties.comments = "Generated via Python 3 and `docx`"
    doc.core_properties.created = datetime.utcnow()
    doc.core_properties.keywords = "résumé cv"

    lb2_style = doc.styles["List Bullet 2"]
    lb2_style.paragraph_format.left_indent = 0
    lb2_style.paragraph_format.space_after = 0

    bt_style = doc.styles["Body Text"]
    bt_style.paragraph_format.space_after = 0

    sec = doc.sections[-1]
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)

    doc.add_heading(cv.full_name)
    urls = " · ".join(l.url for l in cv.links.all())
    p = f"{cv.email} · {cv.phone}"
    if urls:
        p = f"{p}\n{urls}"

    doc.add_paragraph(p)
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(cv.summary)

    doc.add_heading("Skills", level=2)

    skills = [(sk.category, sk.values) for sk in cv.skills.all()]
    if skills:
        if 0:
            mx = 1 + max([len(s[0]) for s in skills])
            p = doc.add_paragraph()
            for i, skill in enumerate(skills):
                cat, values = skill
                p.add_run("{}{:{}}".format("\n" if i else "", "{}:".format(cat), mx)).bold = True
                p.add_run(f"\n{values}")
        else:
            table = doc.add_table(rows=0, cols=0)
            table.autofit = False
            table.allow_autofit = False
            table.add_column(Inches(1.25))
            table.add_column(Inches(6.0))
            for i, (cat, values) in enumerate(skills):
                row = table.add_row()
                cells = row.cells
                cells[0].width = Inches(1.25)
                cells[0].text = f"{cat}:"
                cells[1].width = Inches(6.0)
                cells[1].text = values

    doc.add_heading("Work Experience", level=2)
    for org in cv.organizations.all():
        p = doc.add_paragraph()
        p.style = bt_style
        r = p.add_run(f"{org.name} ")
        r.bold = True
        r.font.size = Pt(13)
        p.add_run(f"· {org.location}\n").italic = True

        for i, pos in enumerate(org.positions.all()):
            if i:
                p = doc.add_paragraph()
                p.style = bt_style

            p.add_run(f"{pos.title} ").bold = True
            p.add_run(f"· {pos.started} - {pos.ended}").italic = True

            desc = pos.summary
            if desc:
                p.add_run(f"\n{desc}").italic = True

            achs = list(pos.iter_achievements)
            if achs:
                for ach in pos.iter_achievements:
                    p = doc.add_paragraph(ach)
                    p.style = lb2_style

                p.style = "List Bullet"
            else:
                p.add_run("\n")

    if cv.educations.count():
        doc.add_heading("Education", level=2)
        p = doc.add_paragraph()
        for i, edu in enumerate(cv.educations.all()):
            deg = edu.result
            dis = edu.focus
            p.add_run(
                "{}{}, {}{}{}".format(
                    "\n" if i else "",
                    edu.institution,
                    edu.location,
                    f", {deg}" if deg else "",
                    f", {dis}" if dis else "",
                )
            )

    p = doc.add_paragraph()
    p.add_run("Last Modified: {}".format(cv.date_updated)).italic = True
    doc.save(stream)
